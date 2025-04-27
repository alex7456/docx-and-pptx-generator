from docx import Document
from generator.summarizer import smart_conclusion

def generate_report(title, intro, sections, conclusion):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_heading("Введение", level=1)
    doc.add_paragraph(intro)
    doc.add_heading("Основная часть", level=1)
    for sec_title, sec_text in sections:
        doc.add_heading(sec_title, level=2)
        doc.add_paragraph(sec_text)

    doc.add_heading("Заключение", level=1)
    final_thought = smart_conclusion(title, sections)
    doc.add_paragraph(final_thought)

    filename = f"{title}_report.docx"
    doc.save(filename)
